#!/usr/bin/env python3
"""
Index Builders for Skippy Multi-Coach System

This module provides specialized index builders for different types of data sources
and coach requirements.
"""

import os
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import json
import sqlite3
from datetime import datetime

# Document processing imports
import PyPDF2
import docx
from docx import Document as DocxDocument
import pandas as pd

# Vector database imports
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Source code analysis imports
import ast
import re
from collections import defaultdict

# =============================================================================
# BASE INDEX BUILDER
# =============================================================================

class BaseIndexBuilder(ABC):
    """Abstract base class for index builders"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.embeddings = self._init_embeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        logging.info(f"{self.__class__.__name__} initialized")
    
    def _init_embeddings(self) -> AzureOpenAIEmbeddings:
        """Initialize embeddings"""
        return AzureOpenAIEmbeddings(
            base_url=self.config.get("embedding_endpoint", ""),
            openai_api_key=self.config.get("embedding_api_key", ""),
            api_version=self.config.get("embedding_api_version", "2022-12-01"),
            model=self.config.get("embedding_deployment", "text-embedding-ada-002"),
            openai_api_type="azure"
        )
    
    @abstractmethod
    def build_index(self, data_path: str, output_path: str) -> bool:
        """Build index for the specific data type"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        pass
    
    def _create_chroma_db(self, output_path: str, documents: List[Document]) -> bool:
        """Create ChromaDB from documents"""
        try:
            # Create output directory
            Path(output_path).mkdir(parents=True, exist_ok=True)
            
            # Create ChromaDB
            vector_db = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=output_path
            )
            
            # Persist the database
            vector_db.persist()
            
            logging.info(f"ChromaDB created successfully at {output_path}")
            return True
            
        except Exception as e:
            logging.error(f"Failed to create ChromaDB: {e}")
            return False

# =============================================================================
# DOCUMENT INDEX BUILDER
# =============================================================================

class DocumentIndexBuilder(BaseIndexBuilder):
    """Index builder for PDF and Word documents"""
    
    def get_supported_extensions(self) -> List[str]:
        """Get supported document extensions"""
        return ['.pdf', '.docx', '.doc']
    
    def build_index(self, data_path: str, output_path: str) -> bool:
        """Build index from documents"""
        try:
            documents = []
            data_dir = Path(data_path)
            
            if not data_dir.exists():
                logging.error(f"Data directory not found: {data_path}")
                return False
            
            # Process all supported files
            for file_path in data_dir.rglob("*"):
                if file_path.suffix.lower() in self.get_supported_extensions():
                    logging.info(f"Processing document: {file_path}")
                    
                    # Extract text based on file type
                    if file_path.suffix.lower() == '.pdf':
                        text = self._extract_pdf_text(file_path)
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        text = self._extract_word_text(file_path)
                    else:
                        continue
                    
                    if text:
                        # Split text into chunks
                        chunks = self.text_splitter.split_text(text)
                        
                        # Create documents
                        for i, chunk in enumerate(chunks):
                            doc = Document(
                                page_content=chunk,
                                metadata={
                                    'source': str(file_path),
                                    'chunk_id': i,
                                    'file_type': file_path.suffix.lower(),
                                    'file_name': file_path.name,
                                    'file_size': file_path.stat().st_size,
                                    'created_at': datetime.now().isoformat()
                                }
                            )
                            documents.append(doc)
            
            if not documents:
                logging.warning("No documents found to index")
                return False
            
            # Create ChromaDB
            return self._create_chroma_db(output_path, documents)
            
        except Exception as e:
            logging.error(f"Document index building failed: {e}")
            return False
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                return text.strip()
                
        except Exception as e:
            logging.error(f"Failed to extract PDF text from {file_path}: {e}")
            return ""
    
    def _extract_word_text(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"Failed to extract Word text from {file_path}: {e}")
            return ""

# =============================================================================
# SOURCE CODE INDEX BUILDER
# =============================================================================

class SourceCodeIndexBuilder(BaseIndexBuilder):
    """Index builder for source code analysis"""
    
    def get_supported_extensions(self) -> List[str]:
        """Get supported source code extensions"""
        return ['.py', '.java', '.js', '.ts', '.cs', '.cpp', '.c', '.h', '.hpp', '.sql', '.xml', '.json', '.yaml', '.yml']
    
    def build_index(self, data_path: str, output_path: str) -> bool:
        """Build index from source code"""
        try:
            documents = []
            data_dir = Path(data_path)
            
            if not data_dir.exists():
                logging.error(f"Source code directory not found: {data_path}")
                return False
            
            # Analyze codebase structure
            codebase_analysis = self._analyze_codebase_structure(data_dir)
            
            # Process source files
            for file_path in data_dir.rglob("*"):
                if file_path.suffix.lower() in self.get_supported_extensions():
                    logging.info(f"Processing source file: {file_path}")
                    
                    # Extract code information
                    code_info = self._extract_code_info(file_path)
                    
                    if code_info:
                        # Create documents for different aspects
                        documents.extend(self._create_code_documents(file_path, code_info, codebase_analysis))
            
            if not documents:
                logging.warning("No source code found to index")
                return False
            
            # Create ChromaDB
            return self._create_chroma_db(output_path, documents)
            
        except Exception as e:
            logging.error(f"Source code index building failed: {e}")
            return False
    
    def _analyze_codebase_structure(self, code_dir: Path) -> Dict[str, Any]:
        """Analyze overall codebase structure"""
        analysis = {
            'total_files': 0,
            'total_lines': 0,
            'languages': defaultdict(int),
            'file_types': defaultdict(int),
            'directories': [],
            'dependencies': set(),
            'entry_points': [],
            'test_files': [],
            'config_files': []
        }
        
        for file_path in code_dir.rglob("*"):
            if file_path.is_file():
                analysis['total_files'] += 1
                analysis['file_types'][file_path.suffix.lower()] += 1
                
                # Count lines
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = len(f.readlines())
                        analysis['total_lines'] += lines
                except:
                    pass
                
                # Categorize files
                if 'test' in file_path.name.lower():
                    analysis['test_files'].append(str(file_path))
                elif file_path.name.lower() in ['main.py', 'app.py', 'index.js', 'main.java']:
                    analysis['entry_points'].append(str(file_path))
                elif file_path.suffix.lower() in ['.json', '.yaml', '.yml', '.xml', '.properties']:
                    analysis['config_files'].append(str(file_path))
        
        return analysis
    
    def _extract_code_info(self, file_path: Path) -> Dict[str, Any]:
        """Extract information from source code file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            code_info = {
                'content': content,
                'lines': len(content.splitlines()),
                'classes': [],
                'functions': [],
                'imports': [],
                'comments': [],
                'complexity': 'low'
            }
            
            # Language-specific analysis
            if file_path.suffix.lower() == '.py':
                code_info.update(self._analyze_python_code(content))
            elif file_path.suffix.lower() in ['.js', '.ts']:
                code_info.update(self._analyze_javascript_code(content))
            elif file_path.suffix.lower() == '.java':
                code_info.update(self._analyze_java_code(content))
            elif file_path.suffix.lower() == '.sql':
                code_info.update(self._analyze_sql_code(content))
            
            return code_info
            
        except Exception as e:
            logging.error(f"Failed to extract code info from {file_path}: {e}")
            return None
    
    def _analyze_python_code(self, content: str) -> Dict[str, Any]:
        """Analyze Python code"""
        try:
            tree = ast.parse(content)
            
            classes = []
            functions = []
            imports = []
            
            for node in ast.walk(tree):
                if isinstance(node, ast.ClassDef):
                    classes.append({
                        'name': node.name,
                        'line': node.lineno,
                        'methods': [n.name for n in node.body if isinstance(n, ast.FunctionDef)]
                    })
                elif isinstance(node, ast.FunctionDef):
                    functions.append({
                        'name': node.name,
                        'line': node.lineno,
                        'args': [arg.arg for arg in node.args.args]
                    })
                elif isinstance(node, (ast.Import, ast.ImportFrom)):
                    if isinstance(node, ast.Import):
                        imports.extend([alias.name for alias in node.names])
                    else:
                        imports.append(f"from {node.module} import {', '.join([alias.name for alias in node.names])}")
            
            return {
                'classes': classes,
                'functions': functions,
                'imports': imports,
                'complexity': 'high' if len(classes) > 10 or len(functions) > 20 else 'medium'
            }
            
        except:
            return {'classes': [], 'functions': [], 'imports': [], 'complexity': 'low'}
    
    def _analyze_javascript_code(self, content: str) -> Dict[str, Any]:
        """Analyze JavaScript/TypeScript code"""
        classes = []
        functions = []
        imports = []
        
        # Simple regex-based analysis
        class_pattern = r'class\s+(\w+)'
        function_pattern = r'function\s+(\w+)|const\s+(\w+)\s*=\s*\(|(\w+)\s*:\s*function'
        import_pattern = r'import\s+.*?from\s+[\'"]([^\'"]+)[\'"]'
        
        classes.extend(re.findall(class_pattern, content))
        functions.extend([f for f in re.findall(function_pattern, content) if f])
        imports.extend(re.findall(import_pattern, content))
        
        return {
            'classes': [{'name': name, 'line': 0} for name in classes],
            'functions': [{'name': name, 'line': 0} for name in functions],
            'imports': imports,
            'complexity': 'high' if len(classes) > 10 or len(functions) > 20 else 'medium'
        }
    
    def _analyze_java_code(self, content: str) -> Dict[str, Any]:
        """Analyze Java code"""
        classes = []
        functions = []
        imports = []
        
        # Simple regex-based analysis
        class_pattern = r'class\s+(\w+)'
        method_pattern = r'(public|private|protected)?\s*(static)?\s*\w+\s+(\w+)\s*\('
        import_pattern = r'import\s+([^;]+);'
        
        classes.extend(re.findall(class_pattern, content))
        functions.extend([match[2] for match in re.findall(method_pattern, content) if match[2]])
        imports.extend(re.findall(import_pattern, content))
        
        return {
            'classes': [{'name': name, 'line': 0} for name in classes],
            'functions': [{'name': name, 'line': 0} for name in functions],
            'imports': imports,
            'complexity': 'high' if len(classes) > 10 or len(functions) > 20 else 'medium'
        }
    
    def _analyze_sql_code(self, content: str) -> Dict[str, Any]:
        """Analyze SQL code"""
        tables = []
        queries = []
        
        # Extract table names
        table_pattern = r'FROM\s+(\w+)|JOIN\s+(\w+)|INTO\s+(\w+)|UPDATE\s+(\w+)'
        tables.extend([t for t in re.findall(table_pattern, content, re.IGNORECASE) if t])
        
        # Extract query types
        query_pattern = r'(SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER)'
        queries.extend(re.findall(query_pattern, content, re.IGNORECASE))
        
        return {
            'classes': [],
            'functions': [],
            'imports': [],
            'tables': list(set(tables)),
            'queries': list(set(queries)),
            'complexity': 'high' if len(tables) > 20 else 'medium'
        }
    
    def _create_code_documents(self, file_path: Path, code_info: Dict[str, Any], codebase_analysis: Dict[str, Any]) -> List[Document]:
        """Create documents from code information"""
        documents = []
        
        # Document 1: File overview
        overview_content = f"""
File: {file_path.name}
Path: {file_path}
Lines: {code_info['lines']}
Complexity: {code_info['complexity']}

Classes: {', '.join([c['name'] for c in code_info['classes']])}
Functions: {', '.join([f['name'] for f in code_info['functions']])}
Imports: {', '.join(code_info['imports'])}

Content:
{code_info['content'][:2000]}...
"""
        
        doc1 = Document(
            page_content=overview_content,
            metadata={
                'source': str(file_path),
                'type': 'file_overview',
                'file_name': file_path.name,
                'file_type': file_path.suffix.lower(),
                'lines': code_info['lines'],
                'complexity': code_info['complexity'],
                'classes_count': len(code_info['classes']),
                'functions_count': len(code_info['functions']),
                'created_at': datetime.now().isoformat()
            }
        )
        documents.append(doc1)
        
        # Document 2: Classes and functions
        if code_info['classes'] or code_info['functions']:
            structure_content = f"""
File Structure: {file_path.name}

Classes:
{chr(10).join([f"- {c['name']} (line {c['line']})" for c in code_info['classes']])}

Functions:
{chr(10).join([f"- {f['name']} (line {f['line']})" for f in code_info['functions']])}

Imports:
{chr(10).join([f"- {imp}" for imp in code_info['imports']])}
"""
            
            doc2 = Document(
                page_content=structure_content,
                metadata={
                    'source': str(file_path),
                    'type': 'code_structure',
                    'file_name': file_path.name,
                    'file_type': file_path.suffix.lower(),
                    'created_at': datetime.now().isoformat()
                }
            )
            documents.append(doc2)
        
        return documents

# =============================================================================
# DATABASE SCHEMA INDEX BUILDER
# =============================================================================

class DatabaseSchemaIndexBuilder(BaseIndexBuilder):
    """Index builder for database schemas"""
    
    def get_supported_extensions(self) -> List[str]:
        """Get supported database schema extensions"""
        return ['.sql', '.ddl', '.schema']
    
    def build_index(self, data_path: str, output_path: str) -> bool:
        """Build index from database schema files"""
        try:
            documents = []
            data_dir = Path(data_path)
            
            if not data_dir.exists():
                logging.error(f"Database schema directory not found: {data_path}")
                return False
            
            # Process schema files
            for file_path in data_dir.rglob("*"):
                if file_path.suffix.lower() in self.get_supported_extensions():
                    logging.info(f"Processing schema file: {file_path}")
                    
                    # Extract schema information
                    schema_info = self._extract_schema_info(file_path)
                    
                    if schema_info:
                        documents.extend(self._create_schema_documents(file_path, schema_info))
            
            if not documents:
                logging.warning("No database schema found to index")
                return False
            
            # Create ChromaDB
            return self._create_chroma_db(output_path, documents)
            
        except Exception as e:
            logging.error(f"Database schema index building failed: {e}")
            return False
    
    def _extract_schema_info(self, file_path: Path) -> Dict[str, Any]:
        """Extract schema information from file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Parse SQL schema
            tables = []
            indexes = []
            constraints = []
            
            # Extract table definitions
            table_pattern = r'CREATE\s+TABLE\s+(\w+)\s*\((.*?)\)'
            for match in re.finditer(table_pattern, content, re.IGNORECASE | re.DOTALL):
                table_name = match.group(1)
                table_def = match.group(2)
                
                # Extract columns
                columns = []
                column_pattern = r'(\w+)\s+(\w+)(?:\s+([^,\n]+))?'
                for col_match in re.finditer(column_pattern, table_def):
                    columns.append({
                        'name': col_match.group(1),
                        'type': col_match.group(2),
                        'constraints': col_match.group(3) or ''
                    })
                
                tables.append({
                    'name': table_name,
                    'columns': columns,
                    'definition': match.group(0)
                })
            
            # Extract indexes
            index_pattern = r'CREATE\s+(?:UNIQUE\s+)?INDEX\s+(\w+)\s+ON\s+(\w+)\s*\((.*?)\)'
            for match in re.finditer(index_pattern, content, re.IGNORECASE):
                indexes.append({
                    'name': match.group(1),
                    'table': match.group(2),
                    'columns': match.group(3)
                })
            
            return {
                'content': content,
                'tables': tables,
                'indexes': indexes,
                'constraints': constraints
            }
            
        except Exception as e:
            logging.error(f"Failed to extract schema info from {file_path}: {e}")
            return None
    
    def _create_schema_documents(self, file_path: Path, schema_info: Dict[str, Any]) -> List[Document]:
        """Create documents from schema information"""
        documents = []
        
        # Document 1: Schema overview
        overview_content = f"""
Database Schema: {file_path.name}

Tables ({len(schema_info['tables'])}):
{chr(10).join([f"- {t['name']} ({len(t['columns'])} columns)" for t in schema_info['tables']])}

Indexes ({len(schema_info['indexes'])}):
{chr(10).join([f"- {i['name']} on {i['table']}" for i in schema_info['indexes']])}

Schema Content:
{schema_info['content'][:2000]}...
"""
        
        doc1 = Document(
            page_content=overview_content,
            metadata={
                'source': str(file_path),
                'type': 'schema_overview',
                'file_name': file_path.name,
                'tables_count': len(schema_info['tables']),
                'indexes_count': len(schema_info['indexes']),
                'created_at': datetime.now().isoformat()
            }
        )
        documents.append(doc1)
        
        # Document 2: Individual tables
        for table in schema_info['tables']:
            table_content = f"""
Table: {table['name']}

Columns:
{chr(10).join([f"- {col['name']}: {col['type']} {col['constraints']}" for col in table['columns']])}

Definition:
{table['definition']}
"""
            
            doc2 = Document(
                page_content=table_content,
                metadata={
                    'source': str(file_path),
                    'type': 'table_definition',
                    'table_name': table['name'],
                    'columns_count': len(table['columns']),
                    'created_at': datetime.now().isoformat()
                }
            )
            documents.append(doc2)
        
        return documents

# =============================================================================
# INDEX BUILDER FACTORY
# =============================================================================

class IndexBuilderFactory:
    """Factory for creating index builders"""
    
    @staticmethod
    def create_builder(builder_type: str, config: Dict[str, Any]) -> BaseIndexBuilder:
        """Create an index builder of the specified type"""
        if builder_type == 'document':
            return DocumentIndexBuilder(config)
        elif builder_type == 'source_code':
            return SourceCodeIndexBuilder(config)
        elif builder_type == 'database_schema':
            return DatabaseSchemaIndexBuilder(config)
        else:
            raise ValueError(f"Unknown builder type: {builder_type}")
    
    @staticmethod
    def build_all_indexes(config: Dict[str, Any], data_paths: Dict[str, str], output_paths: Dict[str, str]) -> Dict[str, bool]:
        """Build all indexes for different data types"""
        results = {}
        
        for data_type, data_path in data_paths.items():
            if data_type in output_paths:
                try:
                    builder = IndexBuilderFactory.create_builder(data_type, config)
                    result = builder.build_index(data_path, output_paths[data_type])
                    results[data_type] = result
                    logging.info(f"Index building for {data_type}: {'Success' if result else 'Failed'}")
                except Exception as e:
                    logging.error(f"Failed to build index for {data_type}: {e}")
                    results[data_type] = False
        
        return results
